import docx
from typing import Dict, Any
from .base_converter import BaseConverter
from ..utils.medical_terms import extract_medical_terms
from ..utils.table_extractor import extract_tables

class DocxConverter(BaseConverter):
    """Converter for DOCX documents"""

    def convert(self, file_path: str, **kwargs) -> str:
        """
        Convert DOCX file to text

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        doc = docx.Document(file_path)
        
        # Extract text with formatting
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Extract tables
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(' | '.join(row_data))
            if rows:
                table_texts.append('\n'.join(rows))
        
        # Combine all text
        all_text = '\n\n'.join(paragraphs)
        if table_texts:
            all_text += '\n\nТаблицы:\n' + '\n\n'.join(table_texts)
        
        return all_text
    
    def get_supported_formats(self) -> list:
        return ['.docx']  # Removed .doc since it's handled separately